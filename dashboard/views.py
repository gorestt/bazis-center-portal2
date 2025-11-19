
import os
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator
from django.http import JsonResponse, FileResponse, HttpResponseForbidden
from django.utils import timezone
from django.conf import settings
from django.contrib.auth.models import User

from .models import OrderQueue, KPIRecord, Incident, Shift, Document, Report
from .forms import OrderForm, DocumentForm, ReportForm

def get_role(user):
    profile = getattr(user, 'profile', None)
    if profile:
        return profile.role
    if user.is_superuser:
        return 'admin'
    return 'client'



def ensure_sample_data():
    # пользователи для связей
    admin = User.objects.filter(username='admin').first()
    manager = User.objects.filter(username='manager').first()
    client = User.objects.filter(username='client').first()

    # Очередь заявок — наполнение до 30 записей за последние 30 дней
    count_orders = OrderQueue.objects.count()
    if count_orders < 30 and client:
        base = timezone.now()
        to_create = 30 - count_orders
        for i in range(to_create):
            day = base - timezone.timedelta(days=i)
            status = 'new'
            if i % 3 == 1:
                status = 'in_progress'
            elif i % 3 == 2:
                status = 'done'
            priority = 'medium'
            if i % 5 == 0:
                priority = 'high'
            elif i % 4 == 0:
                priority = 'low'
            OrderQueue.objects.create(
                title=f'Инцидент с сервисом №{count_orders + i + 1}',
                description='Заявка сформирована для демонстрации загрузки очереди за месяц.',
                initiator=client,
                executor=manager or admin,
                status=status,
                priority=priority,
                created_at=day,
                sla_deadline=day + timezone.timedelta(hours=4),
            )

    # KPI — формируются только один раз, за 30 дней
    if KPIRecord.objects.count() == 0:
        base = timezone.now()
        for i in range(30):
            day = base - timezone.timedelta(days=i)
            KPIRecord.objects.create(
                metric='Среднее время решения',
                value=4.5 - i * 0.05,
                timestamp=day,
                service_name='Портал клиентов',
            )
            KPIRecord.objects.create(
                metric='Доступность сервиса',
                value=98.5 + i * 0.03,
                timestamp=day,
                service_name='Система биллинга',
            )

    # Инциденты — наполнение до 30 записей за последние 30 дней
    count_inc = Incident.objects.count()
    if count_inc < 30:
        base = timezone.now()
        orders = list(OrderQueue.objects.all())
        to_create = 30 - count_inc
        for i in range(to_create):
            day = base - timezone.timedelta(days=i)
            crit = 'medium'
            if i % 4 == 0:
                crit = 'high'
            elif i % 3 == 0:
                crit = 'low'
            status = 'В работе'
            closed_at = None
            if i % 2 == 0:
                status = 'Закрыт'
                closed_at = day + timezone.timedelta(hours=2)
            related = orders[(count_inc + i) % len(orders)] if orders else None
            Incident.objects.create(
                title=f'Инцидент производственного контура №{count_inc + i + 1}',
                description='Инцидент создан для заполнения статистики по отказам и восстановлениям за месяц.',
                status=status,
                criticality=crit,
                detected_at=day,
                closed_at=closed_at,
                related_order=related,
            )

    # Смены — наполнение до 30 записей вперёд от текущей даты
    count_shifts = Shift.objects.count()
    if count_shifts < 30 and (admin or manager):
        base_date = timezone.now().date()
        to_create = 30 - count_shifts
        for i in range(to_create):
            Shift.objects.create(
                employee=manager or admin,
                date=base_date + timezone.timedelta(days=i),
                shift='day',
                comment='Плановая дневная смена дежурного инженера.',
                phone='+7 (900) 000-00-01',
            )
        for i in range(0, min(10, to_create), 3):
            Shift.objects.create(
                employee=admin or manager,
                date=base_date + timezone.timedelta(days=i),
                shift='night',
                comment='Ночная смена дежурного персонала.',
                phone='+7 (900) 000-00-02',
            )

    # Документы
    docs_dir = settings.MEDIA_ROOT / 'docs'
    os.makedirs(docs_dir, exist_ok=True)
    if Document.objects.count() == 0:
        try:
            def make_doc(filename, title, body_text):
                full_path = docs_dir / filename
                if not full_path.exists():
                    with open(full_path, 'w', encoding='utf8') as f:
                        f.write(body_text)

            make_doc('reglament_incidents.docx', 'Регламент обработки инцидентов',
                     'Документ описывает порядок регистрации, классификации и эскалации инцидентов.')
            make_doc('reglament_shifts.docx', 'Регламент организации смен',
                     'Документ фиксирует правила формирования графика смен и обязанности дежурного персонала.')
            make_doc('instruction_portal.docx', 'Инструкция пользователя портала',
                     'Инструкция по работе с порталом и отслеживанию статусов обращений.')
        except Exception:
            def make_txt(filename, body_text):
                full_path = docs_dir / filename
                with open(full_path, 'w', encoding='utf8') as f:
                    f.write(body_text)
            make_txt('reglament_incidents.txt', 'Регламент обработки инцидентов.')
            make_txt('reglament_shifts.txt', 'Регламент организации смен.')
            make_txt('instruction_portal.txt', 'Инструкция пользователя портала.')

        for fname, title, desc in [
            ('reglament_incidents.docx', 'Регламент обработки инцидентов', 'Порядок регистрации и сопровождения инцидентов.'),
            ('reglament_shifts.docx', 'Регламент организации смен', 'Описание процедуры планирования смен и ответственности.'),
            ('instruction_portal.docx', 'Инструкция пользователя портала', 'Руководство по работе с порталом для клиентов.'),
        ]:
            path_docx = docs_dir / fname
            if path_docx.exists():
                rel = path_docx.relative_to(settings.MEDIA_ROOT)
                Document.objects.get_or_create(
                    title=title,
                    defaults={
                        'description': desc,
                        'access': 'public',
                        'file': str(rel),
                    },
                )
            else:
                alt = fname.replace('.docx', '.txt')
                path_txt = docs_dir / alt
                if path_txt.exists():
                    rel = path_txt.relative_to(settings.MEDIA_ROOT)
                    Document.objects.get_or_create(
                        title=title,
                        defaults={
                            'description': desc,
                            'access': 'public',
                            'file': str(rel),
                        },
                    )
@login_required
def home(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role == 'client':
        return redirect('dashboard:client_home')
    open_orders = OrderQueue.objects.filter(status__in=['new','in_progress']).count()
    incidents_open = Incident.objects.exclude(status='Закрыт').count()
    kpi_count = KPIRecord.objects.count()
    return render(request, 'dashboard/home.html', {
        'open_orders': open_orders,
        'incidents_open': incidents_open,
        'kpi_count': kpi_count,
    })

@login_required
def queue_list(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager']:
        return HttpResponseForbidden('Доступ запрещён.')
    qs = OrderQueue.objects.select_related('initiator','executor').order_by('-created_at')
    status = request.GET.get('status')
    priority = request.GET.get('priority')
    if status:
        qs = qs.filter(status=status)
    if priority:
        qs = qs.filter(priority=priority)
    paginator = Paginator(qs, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    return render(request, 'dashboard/queue_list.html', {
        'page_obj': page_obj,
        'status': status,
        'priority': priority,
    })

@login_required
def queue_create(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager', 'client']:
        return HttpResponseForbidden('Доступ запрещён.')
    if request.method == 'POST':
        form = OrderForm(request.POST)
        if form.is_valid():
            order = form.save(commit=False)
            order.initiator = request.user
            order.save()
            return redirect('dashboard:queue_list')
    else:
        form = OrderForm()
    return render(request, 'dashboard/queue_form.html', {'form': form, 'mode': 'create'})

@login_required
def queue_edit(request, pk):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager']:
        return HttpResponseForbidden('Доступ запрещён.')
    order = get_object_or_404(OrderQueue, pk=pk)
    if request.method == 'POST':
        form = OrderForm(request.POST, instance=order)
        if form.is_valid():
            form.save()
            return redirect('dashboard:queue_list')
    else:
        form = OrderForm(instance=order)
    return render(request, 'dashboard/queue_form.html', {'form': form, 'mode': 'edit', 'order': order})

@login_required
def queue_detail(request, pk):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager', 'client']:
        return HttpResponseForbidden('Доступ запрещён.')
    order = get_object_or_404(OrderQueue, pk=pk)
    # клиент видит только свои заявки
    if role == 'client' and order.initiator != request.user:
        return HttpResponseForbidden('Доступ запрещён.')
    return render(request, 'dashboard/queue_detail.html', {'order': order})

@login_required
def kpi_dashboard(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager']:
        return HttpResponseForbidden('Доступ запрещён.')
    last_30 = timezone.now() - timezone.timedelta(days=30)
    kpi = KPIRecord.objects.filter(timestamp__gte=last_30)
    series = {}
    for rec in kpi:
        series.setdefault(rec.metric, []).append({
            'timestamp': rec.timestamp.isoformat(),
            'value': rec.value,
        })
    return render(request, 'dashboard/kpi_dashboard.html', {
        'kpi': kpi,
        'series_json': series,
    })

@login_required
def incidents_list(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager']:
        return HttpResponseForbidden('Доступ запрещён.')
    items = Incident.objects.order_by('-detected_at')
    paginator = Paginator(items, 20)
    page_obj = paginator.get_page(request.GET.get('page'))
    return render(request, 'dashboard/incidents_list.html', {'page_obj': page_obj})

@login_required
def shifts_list(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager']:
        return HttpResponseForbidden('Доступ запрещён.')
    items = Shift.objects.select_related('employee').order_by('date')
    return render(request, 'dashboard/shifts_list.html', {'items': items})

@login_required

def reports_panel(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager']:
        return HttpResponseForbidden('Доступ запрещён.')

    if request.method == 'POST':
        form = ReportForm(request.POST)
        if form.is_valid():
            # вычисляем период автоматически исходя из типа
            report_type = form.cleaned_data['report_type']
            period_to = form.cleaned_data['period_to']
            if report_type == 'daily':
                period_from = period_to
            elif report_type == 'weekly':
                period_from = period_to - timedelta(days=6)
            else:  # monthly
                period_from = period_to.replace(day=1)

            report = Report(
                report_type=report_type,
                period_from=period_from,
                period_to=period_to,
                author=request.user,
            )

            # собираем агрегаты для отчёта
            orders = OrderQueue.objects.filter(created_at__date__gte=period_from,
                                               created_at__date__lte=period_to)
            incidents = Incident.objects.filter(detected_at__date__gte=period_from,
                                               detected_at__date__lte=period_to)
            kpi_records = KPIRecord.objects.filter(timestamp__date__gte=period_from,
                                                   timestamp__date__lte=period_to)

            try:
                from docx import Document as DocxDocument
                doc = DocxDocument()
                doc.add_heading('Сводный отчёт по сервисам и заявкам', level=1)
                doc.add_paragraph(f"Тип отчёта: {report.get_report_type_display()}")
                doc.add_paragraph(f"Период: {report.period_from} – {report.period_to}")
                doc.add_paragraph(' ')

                # Блок по заявкам
                doc.add_heading('Заявки службы поддержки', level=2)
                total_orders = orders.count()
                new_count = orders.filter(status='new').count()
                prog_count = orders.filter(status='in_progress').count()
                done_count = orders.filter(status='done').count()
                doc.add_paragraph(
                    f"Всего заявок за период: {total_orders} "
                    f"(новые: {new_count}, в работе: {prog_count}, закрытые: {done_count})."
                )

                # Блок по инцидентам
                doc.add_heading('Инциденты', level=2)
                total_inc = incidents.count()
                high_inc = incidents.filter(criticality='high').count()
                doc.add_paragraph(
                    f"Всего инцидентов: {total_inc}, из них критичных: {high_inc}."
                )

                # Блок по KPI
                doc.add_heading('KPI сервисов', level=2)
                if kpi_records.exists():
                    metrics = {}
                    for rec in kpi_records:
                        metrics.setdefault(rec.metric, []).append(rec.value)
                    table = doc.add_table(rows=1, cols=3)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Метрика'
                    hdr_cells[1].text = 'Среднее значение'
                    hdr_cells[2].text = 'Кол-во записей'
                    for metric, values in metrics.items():
                        row_cells = table.add_row().cells
                        row_cells[0].text = metric
                        row_cells[1].text = f"{sum(values) / len(values):.2f}"
                        row_cells[2].text = str(len(values))
                else:
                    doc.add_paragraph('Записей KPI за выбранный период не найдено.')

                # сохраняем docx
                path_dir = settings.MEDIA_ROOT / 'reports'
                os.makedirs(path_dir, exist_ok=True)
                filename = f"report_{report.report_type}_{report.period_from}_{report.period_to}.docx"
                full_path = path_dir / filename
                doc.save(full_path)
                report.file.name = f"reports/{filename}"
            except Exception:
                # резерв: простой текстовый отчёт
                path_dir = settings.MEDIA_ROOT / 'reports'
                os.makedirs(path_dir, exist_ok=True)
                filename = f"report_{report.report_type}_{report.period_from}_{report.period_to}.txt"
                full_path = path_dir / filename
                with open(full_path, 'w', encoding='utf8') as f:
                    f.write('Сводный отчёт по сервисам и заявкам\n')
                    f.write(f"Тип отчёта: {report.get_report_type_display()}\n")
                    f.write(f"Период: {report.period_from} – {report.period_to}\n")
                    f.write(f"Всего заявок: {orders.count()}\n")
                    f.write(f"Всего инцидентов: {incidents.count()}\n")
                    f.write(f"KPI записей: {kpi_records.count()}\n")
                report.file.name = f"reports/{filename}"

            report.save()
            return redirect('dashboard:reports_panel')
    else:
        form = ReportForm()

    reports = Report.objects.order_by('-created_at')
    return render(request, 'dashboard/reports_panel.html', {'reports': reports, 'form': form})
@login_required
def report_download(request, pk):
    role = get_role(request.user)
    if role not in ['admin', 'manager']:
        return HttpResponseForbidden('Доступ запрещён.')
    report = get_object_or_404(Report, pk=pk)
    if not report.file:
        raise HttpResponseForbidden('Файл отчёта не найден.')
    full_path = report.file.path
    return FileResponse(open(full_path, 'rb'), as_attachment=True, filename=os.path.basename(full_path))

@login_required

def docs_manage(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager']:
        return HttpResponseForbidden('Доступ запрещён.')

    docs = Document.objects.all().order_by('title')
    edit_obj = None

    # удаление документа
    if request.method == 'POST' and 'delete_id' in request.POST:
        doc = get_object_or_404(Document, pk=request.POST['delete_id'])
        doc.delete()
        return redirect('dashboard:docs_manage')

    # выбор документа для редактирования (?edit=ID)
    if 'edit' in request.GET:
        edit_obj = get_object_or_404(Document, pk=request.GET['edit'])

    # сохранение изменений
    if request.method == 'POST' and 'save_id' in request.POST:
        edit_obj = get_object_or_404(Document, pk=request.POST['save_id'])
        form = DocumentForm(request.POST, request.FILES, instance=edit_obj)
        if form.is_valid():
            form.save()
            return redirect('dashboard:docs_manage')
    else:
        form = DocumentForm(instance=edit_obj) if edit_obj else None

    return render(request, 'dashboard/docs_manage.html', {
        'docs': docs,
        'form': form,
        'edit_obj': edit_obj,
    })


@login_required
def client_home(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role != 'client':
        return redirect('dashboard:home')
    my_orders = OrderQueue.objects.filter(initiator=request.user).order_by('-created_at')
    return render(request, 'dashboard/client_home.html', {'orders': my_orders})

# API views
@login_required
def queue_api(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager']:
        return HttpResponseForbidden('Доступ запрещён.')
    qs = OrderQueue.objects.order_by('-created_at')
    status = request.GET.get('status')
    priority = request.GET.get('priority')
    if status:
        qs = qs.filter(status=status)
    if priority:
        qs = qs.filter(priority=priority)
    data = [{
        'id': o.id,
        'title': o.title,
        'status': o.status,
        'priority': o.priority,
        'created_at': o.created_at.isoformat(),
    } for o in qs[:200]]
    return JsonResponse({'results': data})

@login_required
def kpi_api(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager']:
        return HttpResponseForbidden('Доступ запрещён.')
    metric = request.GET.get('metric')
    qs = KPIRecord.objects.all()
    if metric:
        qs = qs.filter(metric=metric)
    data = [{
        'metric': r.metric,
        'value': r.value,
        'timestamp': r.timestamp.isoformat(),
        'service_name': r.service_name,
    } for r in qs[:500]]
    return JsonResponse({'results': data})